import pandas as pd
import ollama
import matplotlib.pyplot as plt
import seaborn as sns
from docx import Document
from docx.shared import Inches

# File paths
csv_file_path = "gst_transactions.csv"  # Update with your CSV path
image_file_path = "fraud_visualization.png"
doc_file_path = "GST_Fraud_Detection_Report.docx"

# Step 1: Analyze CSV with Ollama
def analyze_csv_with_ollama(csv_file):
    """Reads the CSV file and uses Ollama (Llama 3.2) for fraud detection analysis."""
    df = pd.read_csv(csv_file)
    data_text = df.head(10).to_string()  # Use first 10 rows for efficiency

    prompt = f"""
    You are an expert in financial fraud detection. Analyze the following GST transaction data and detect potential fraud patterns:

    {data_text}

    Look for:
    - Unusual input tax credit (ITC) claims.
    - Mismatched sales vs. GST paid.
    - Duplicate transactions.
    - Suspicious refunds or tax evasion risks.

    Provide a structured fraud detection report highlighting possible issues.
    """
    response = ollama.chat(model="llama3.2", messages=[{"role": "user", "content": prompt}])
    return response['message']['content']

# Step 2: Visualize Data
def visualize_data(csv_file, image_path):
    """Creates bar and scatter plots for GST transaction data."""
    df = pd.read_csv(csv_file)

    sns.set(style="whitegrid")
    fig, axes = plt.subplots(2, 1, figsize=(10, 8))

    # Barplot: Sales vs GST Paid
    sns.barplot(x="Company Name", y="Sales Amount (₹)", data=df, ax=axes[0], palette="viridis")
    axes[0].set_title("Sales Amount by Company")
    axes[0].set_xticklabels(df["Company Name"], rotation=45)

    # Scatterplot: Input Tax Credit vs Sales Amount
    sns.scatterplot(x="Sales Amount (₹)", y="Input Tax Credit (₹)", hue="Company Name", data=df, ax=axes[1], palette="deep")
    axes[1].set_title("Input Tax Credit vs Sales Amount")

    plt.tight_layout()
    plt.savefig(image_path)
    plt.show()

# Step 3: Save Report to Word Document
def save_report_to_doc(report_text, image_path, output_doc):
    """Saves the fraud detection report and visualization into a Word document."""
    doc = Document()
    doc.add_heading("GST Fraud Detection Report", level=1)

    doc.add_heading("Fraud Analysis Summary", level=2)
    doc.add_paragraph(report_text)

    doc.add_heading("Data Visualizations", level=2)
    doc.add_paragraph("The following chart provides insights into GST transactions and possible anomalies.")
    doc.add_picture(image_path, width=Inches(6))

    doc.save(output_doc)
    print(f"📄 Report saved successfully: {output_doc}")

# Run analysis and generate report
fraud_report = analyze_csv_with_ollama(csv_file_path)
visualize_data(csv_file_path, image_file_path)
save_report_to_doc(fraud_report, image_file_path, doc_file_path)

print("\n✅ Report generation completed!")

